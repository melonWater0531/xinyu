from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
CATALOG = DOCS / "XINYU_FEATURE_DEPLOYMENT_CATALOG.docx"
REPORT = DOCS / "WEEKLY_REPORT_2026-06-25_2026-06-30.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B3A"
MUTED = "65727E"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E8F3EC"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, color=None, italic=None, latin=FONT_LATIN, cjk=FONT_CJK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tblpr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tblpr.find(qn(tag))
        if old is not None:
            tblpr.remove(old)
    tblw = OxmlElement("w:tblW")
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblpr.append(tblw)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(indent_dxa))
    ind.set(qn("w:type"), "dxa")
    tblpr.append(ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[idx]))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8, color=MUTED)


def configure_doc(doc, preset="compact"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.10

    values = {
        "Heading 1": (16, BLUE, 18 if preset == "compact" else 16, 10 if preset == "compact" else 8),
        "Heading 2": (13, BLUE, 14 if preset == "compact" else 12, 7 if preset == "compact" else 6),
        "Heading 3": (12, DARK_BLUE, 10 if preset == "compact" else 8, 5 if preset == "compact" else 4),
    }
    for name, (size, color, before, after) in values.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "compact" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "compact" else 8)
        style.paragraph_format.line_spacing = 1.25 if preset == "compact" else 1.167


def set_running_furniture(section, left, right):
    hp = section.header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run(left)
    set_run_font(r, size=8, bold=True, color=MUTED)
    r = hp.add_run("    " + right)
    set_run_font(r, size=8, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.clear()
    add_page_field(fp)


def add_title(doc, title, subtitle, kicker=None, centered=False):
    if kicker:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(kicker.upper())
        set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, size=29 if centered else 25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=MUTED)


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.18
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    ppr.append(borders)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def catalog_rows():
    rows = []

    def add(group, module, name, desc, deploy, status, limits, verify):
        rows.append((group, module, name, desc, deploy, status, limits, verify))

    add("视觉感知", "视频", "reCamera SSCMA 视频接入", "接收设备 JPEG 帧和检测框。", "配置 RECAMERA_DEVICE_IP；确认 8090 可达；启动 FastAPI；检查 video.connected。", "已部署", "依赖设备网络和 SSCMA 服务。", "/video_feed 连续显示，/api/state 中视频在线。")
    add("视觉感知", "视频", "MJPEG 实时推流", "将最新 JPEG 作为浏览器 MJPEG 输出。", "启动 FastAPI；访问 /video_feed；在 /home 或 /control 查看。", "已部署", "浏览器连接过多会增加编码/网络负载。", "画面持续刷新且断线可恢复。")
    add("视觉感知", "检测", "SSCMA 检测框解析", "解析 cx/cy/w/h/conf/cls 格式目标。", "保持设备模型输出；检查 boxes；在叠层中核对坐标。", "已部署", "依赖设备侧模型类别和坐标约定。", "框位置与实际目标一致。")
    add("视觉感知", "人脸", "人脸检测与跨帧跟踪", "SCRFD、Kalman/ByteTrack 和可选 ArcFace 的多脸跟踪。", "安装推荐人脸依赖；启动感知；用多人进出画面验证 track id。", "已部署", "insightface 为可选依赖；缺失时有降级路径。", "连续帧主目标稳定，短遮挡不立即丢失。")
    add("视觉感知", "人体", "YOLO11 人体姿态", "输出人体框与 COCO-17 关键点。", "确认 models/yolo11n-pose.onnx；启动 FastAPI；检查 pose.persons。", "已部署", "CPU 推理频率受设备性能影响。", "人体数和主要关键点与画面一致。")
    add("视觉感知", "面部", "MediaPipe 面部关键点", "提供眼部、虹膜和面部网格输入。", "确认 face_landmarker.task；安装 mediapipe；观察 face_landmark。", "已部署", "光照、遮挡和脸部尺寸影响关键点。", "有人脸时关键点成功，无人脸时安全降级。")
    add("视觉感知", "情绪", "EmotiEff 8 类情绪识别", "输出情绪、置信度和概率分布。", "确认 EmotiEff ONNX 模型；启动单人感知；查看 emotieff。", "已部署", "表情分类不是心理诊断。", "状态持续更新且置信度范围为 0-1。")
    add("视觉感知", "眼部", "EAR/PERCLOS/眨眼率", "计算眼睛开合、闭眼占比和眨眼指标。", "启用 Face Landmarker；保持面部可见；检查 eye_metrics。", "已部署", "眼镜、眯眼和低分辨率会影响阈值。", "开闭眼变化能反映到指标且无人脸不崩溃。")
    add("视觉感知", "专注", "多证据专注评分", "融合头姿、眼部、稳定性和 gaze，输出 0-100。", "启动单人监测；检查 attention.components 与 score。", "已部署", "是行为趋势指标，不是能力或医疗评价。", "四项组件与权重存在，score 平滑变化。")
    add("视觉感知", "注视", "注视方向估计", "根据虹膜相对眼角偏移输出粗粒度视线趋势。", "确认 478 点虹膜输出；观察 center/left/right/down/away。", "已部署", "不是精确眼动追踪；受眼镜和光照影响。", "无人脸 available=false；方向变化可观察。")
    add("视觉感知", "叠层", "跟踪与控制叠层", "在调试台显示框、轨迹、FSM、遥测和 trace。", "访问 /control；连接 WebSocket；开启对应场景。", "已部署", "用于开发调试，不替代日志和设备遥测。", "叠层不遮挡主画面，状态与 API 一致。")
    add("视觉感知", "快照", "当前帧 JPEG 快照", "通过 API 获取当前画面。", "启动视频；请求 GET /api/snapshot；保存或查看响应。", "已部署", "只提供手动当前帧，未接事件自动存储。", "返回有效 JPEG；无帧时给出明确错误。")

    add("情绪健康", "策略", "主动情绪干预", "180 秒滑动窗口判断持续负面情绪或疲劳/低专注。", "保持 emotieff/attention/eye/gaze 输入；等待窗口；消费 proactive_intervention。", "已部署", "30 分钟策略冷却；不得作为心理诊断。", "不足 3 分钟/低置信不触发；冷却内不重复。")
    add("情绪健康", "日记", "情绪日记与日历", "本地保存每日情绪、专注、文字与回复。", "在 /home 写入今日；检查 localStorage；可调用 /api/reflect。", "已部署", "浏览器本地数据不等于服务端备份。", "刷新后记录保留，选日可查看。")
    add("情绪健康", "专注", "专注记录与计时", "开启监测后累计本次专注时长和状态。", "在 /home 开启专注记录；保持人脸输入；观察时长。", "已部署", "页面计时受刷新和浏览器休眠影响。", "启停状态和时长符合操作。")
    add("情绪健康", "趋势", "周情绪/专注趋势", "汇总七天本地记录并生成趋势视图。", "积累或录入日记；打开趋势页；检查平均专注和主要情绪。", "已部署", "数据来源主要是 localStorage。", "七天数据与图表/汇总一致。")
    add("情绪健康", "健康", "20-20-20 护眼计时", "20 分钟循环后提醒看远处 20 秒。", "在健康页开启护眼提醒；保持页面运行；等待或测试计时。", "已部署", "第一版是前端页面计时。", "到期显示站内提示并尝试发送 eye 通知。")
    add("情绪健康", "健康", "45 分钟久坐计时", "45 分钟循环后提醒活动并展示拉伸。", "在健康页开启久坐提醒；保持页面运行。", "已部署", "不使用姿态或穿戴设备判定真实活动。", "到期显示提醒、拉伸引导和 sit 通知。")
    add("情绪健康", "健康", "饮水记录与目标", "记录杯数、最近饮水时间和目标。", "点击喝了一杯水；调整目标存储；查看计数。", "已部署", "数据仅保存在当前浏览器。", "杯数/时间更新，达标后不再触发喝水提醒。")
    add("情绪健康", "健康", "步数记录", "手工增加步数并显示目标进度。", "在健康页点击记录 500 步；编辑目标。", "已部署", "未接手机或穿戴设备传感器。", "数值和进度条同步更新。")
    add("情绪健康", "健康", "4-7-8 呼吸与快速拉伸", "提供节奏动画和分步动作提示。", "打开健康页；开始呼吸或快速拉伸；按提示完成。", "已部署", "纯前端引导，不提供医疗建议。", "启停正常，提示顺序完整。")
    add("情绪健康", "建议", "健康陪伴建议", "结合当前情绪/专注生成模板或 LLM 建议。", "配置可选 DeepSeek；在 /home 点击生成建议或对话。", "已部署", "LLM 不可用时使用本地 fallback。", "接口失败时仍有温和本地建议。")

    add("手势交互", "框架", "MediaPipe 手势识别 A 版", "识别手势并映射低风险陪伴 intent。", "下载模型至 models/gesture_recognizer.task；重启；检查 gesture。", "已部署（资源待补）", "当前仓库缺模型；不允许进入云台控制。", "模型补齐后 available=true，单帧误识别不 ready。")
    add("手势交互", "Open Palm", "唤起心屿", "张手映射 summon_xinyu，显示“我在听”。", "保持 Open Palm 至少 4 帧；确认 intent_ready；查看聊天区。", "已部署（资源待补）", "不自动录音、不控制云台。", "只触发一次陪伴反馈，3 秒内不重复。")
    add("手势交互", "Closed Fist", "暂停/收起提醒", "握拳映射 pause_or_mute。", "保持 Closed Fist；确认 intent；观察 toast。", "已部署（资源待补）", "当前无 TTS，因此只收起本地提示。", "不影响录音和云台，页面给出已收起反馈。")
    add("手势交互", "Thumb Up", "正向反馈", "点赞映射 feedback_positive。", "保持 Thumb Up；检查 xinyu_gesture_feedback。", "已部署（资源待补）", "只记录本地反馈，不在线训练模型。", "localStorage 记录 intent 与时间。")
    add("手势交互", "Thumb Down", "负向反馈", "点踩映射 feedback_negative。", "保持 Thumb Down；检查本地反馈和 toast。", "已部署（资源待补）", "不自动修改模型或删除内容。", "localStorage 记录正确 intent。")
    add("手势交互", "Victory", "积极瞬间草稿", "剪刀手映射 capture_positive_moment。", "保持 Victory；检查今日快速日记文本框。", "已部署（资源待补）", "只生成草稿，不自动永久保存。", "草稿出现且等待用户确认。")

    add("音频会议", "DOA/VAD", "ReSpeaker USB DOA/VAD", "读取 0-359 度声源方向和语音活动。", "USB 透传设备；设置 RECAMERA_DOA_SOURCE=usb；检查 /api/respeaker/state。", "已部署", "依赖 XVF3800 USB control 权限。", "角度、speech 和 age 持续更新。")
    add("音频会议", "DOA/VAD", "TCP DOA 备用输入", "在 9999 端口接收角度、JSON 或 xvf_host 文本。", "设置 RECAMERA_DOA_SOURCE=tcp；启动；用 send_doa_tcp.py 注入。", "已部署", "备用模式不能控制本机实体 LED。", "注入数据可在状态中观察并具备过期保护。")
    add("音频会议", "灯效", "ReSpeaker DOA LED", "按权威功能模式显示或关闭灯环。", "使用 USB DOA；启动多人/会议功能；观察灯效。", "已部署", "TCP 模式不控制实体灯。", "启停功能时灯效与运行时一致。")
    add("音频会议", "录音", "会议录音与 VAD 分段", "16 kHz 单声道录音并按语音活动切片。", "确认音频设备；设置 RECAMERA_AUDIO_DEVICE；调用 conversation start/stop。", "已部署", "音频设备索引和权限需现场确认。", "会话状态正常，片段可保存/结束。")
    add("音频会议", "ASR", "faster-whisper 本地转写", "对 WAV 片段进行本地语音转写。", "安装 faster-whisper；下载模型；完成录音后触发转写。", "已部署（资源待补）", "依赖额外包、模型、CPU/内存。", "测试音频生成可读文本且失败可降级。")
    add("音频会议", "会议", "会议摘要与日记写入", "汇总转写内容，生成摘要和日记候选。", "完成会议录音/转写；调用 /api/meeting/summarize；确认结果。", "已部署（资源待补）", "完整效果依赖 ASR 和可选 DeepSeek。", "返回摘要，前端能写入会议记录。")
    add("音频会议", "唤醒词", "基础唤醒词模块", "提供可选唤醒检测代码。", "准备模型和音频输入；单独验证 wake_word.py；再设计状态机。", "部分部署", "未接常驻主链路，避免和录音/DOA 抢设备。", "模块单测可用但主页面不宣称常驻唤醒。")

    add("控制安全", "控制面", "单控制平面", "main_phase3.py 是唯一真实硬件命令源。", "启动 main_phase3；连接 EventBus；禁止 FastAPI 直接创建设备控制客户端。", "已部署", "双进程必须使用同一设备 IP。", "代码审计无 FastAPI apply_command 调用。")
    add("控制安全", "FSM", "五态控制状态机", "管理 IDLE/AUDIO_SEARCH/VISION_TRACK/FUSED_TRACK/LOST。", "启动运行时；注入视觉/音频事件；观察 fsm_state。", "已部署", "状态机只决定状态，命令由 Orchestrator 生成。", "事件序列产生预期状态转移。")
    add("控制安全", "跟踪", "单人视觉 yaw/pitch 跟踪", "视觉目标中心驱动双轴对准。", "部署 Node-RED bridge；启动控制；在 /home 或 /control 开启单人。", "已部署", "真实运动需要设备、桥和有效租约。", "目标偏移时命令合理，丢失时进入保护。")
    add("控制安全", "跟踪", "多人声源 yaw 跟随", "DOA 驱动 yaw，pitch 不由音频控制。", "准备 DOA；启动 multi_sound_yaw；保持 heartbeat。", "已部署", "噪声和混响影响 DOA；有新鲜度门控。", "有效语音产生 yaw 命令，过期后停止。")
    add("控制安全", "会议", "会议 yaw 跟随", "会议功能下的独立声源 yaw 模式。", "启动 meeting/yaw；检查 session 与 runtime；结束后 stop。", "已部署", "与其他控制功能互斥。", "功能租约正确，切换时旧会话失效。")
    add("控制安全", "手动", "手动 D-Pad 与回中", "调试台发送相对移动和 home UI Event。", "进入 /control；启动手动会话；按键并保持 heartbeat。", "已部署", "仅供联调；受 SafetyLayer 限制。", "按键对应方向，停止/离页后不再运动。")
    add("控制安全", "会话", "session + 2.5 秒租约", "维持唯一控制权并隔离旧页面。", "启动功能获取 session_id；每秒 heartbeat；测试失联和接管。", "已部署", "租约过期必须 fail closed。", "旧 session 的 heartbeat/stop 无效，失联自动停。")
    add("控制安全", "安全", "SafetyLayer 硬门控", "执行限频、步长、范围和速度限制。", "加载配置；注入越界/高频命令；检查 safety reason。", "已部署", "不能用 UI 绕过。", "危险命令被拒绝或裁剪并有明确原因。")
    add("控制安全", "通信", "EventBus 控制事件", "连接 FastAPI UI Event 与控制运行时。", "启动 main_phase3 EventBus 8765；再启动 FastAPI；请求 runtime。", "已部署", "EventBus 不可达时 UI 控制应显示 unreachable。", "事件被接受且快照能往返。")
    add("控制安全", "硬件", "Node-RED 双轴 bridge", "提供 command、stop、status 到官方电机节点。", "导入 deploy/node_red flow；配置 CAN；部署；访问 status。", "已部署", "设备侧 flow 必须现场部署。", "status 有真实 readback，command/stop 可控。")
    add("控制安全", "停机", "安全退出与 fail closed", "退出、桥不可达、租约失效时停止运动。", "模拟 Ctrl+C、断网、停 bridge 和丢 heartbeat。", "已部署", "真实设备验收必须留安全空间。", "各故障均进入 stop，不保留持续命令。")

    add("后端接口", "服务", "FastAPI 页面与静态资源", "提供 /home、/control、PWA 与静态素材。", "安装 requirements；启动 recamera_fastapi.py；访问入口。", "已部署", "默认端口 8001。", "主要页面和静态资源返回 200。")
    add("后端接口", "状态", "GET /api/state", "返回当前完整状态快照。", "启动服务；curl /api/state；检查新增三个状态块。", "已部署", "状态是瞬时快照。", "JSON 可解析且字段类型稳定。")
    add("后端接口", "状态", "WebSocket /ws", "约每 200 ms 推送状态。", "连接 /ws；持续接收；断开重连。", "已部署", "客户端数量影响广播开销。", "连续收到 JSON 且无泄漏连接。")
    add("后端接口", "设备", "设备地址配置 API", "读取/更新感知设备 IP 并重启视频客户端。", "GET/POST /api/device/config；填写 IP；检查重连。", "已部署", "不创建云台控制客户端。", "视频切换成功且控制边界不变。")
    add("后端接口", "视频", "视频/快照/调试 API", "提供 feed、snapshot 和 debug/video。", "逐个请求接口；断开设备测试错误响应。", "已部署", "依赖最新 JPEG。", "在线返回有效数据，离线返回明确状态。")
    add("后端接口", "控制", "功能启停 API", "单人、多人、会议和手动功能通过 UI Event 启停。", "启动 EventBus；调用 start/stop；检查 runtime。", "已部署", "FastAPI 只发事件。", "active_feature、session、lease 与调用一致。")
    add("后端接口", "录音", "Conversation API", "提供录音开始、停止、保存和调试状态。", "配置音频设备；调用 conversation 路由；检查状态。", "已部署", "需音频硬件/权限。", "状态机和保存结果符合操作。")
    add("后端接口", "健康", "健康检查与运行时遥测", "返回服务、视频、DOA、控制等健康信息。", "访问 /api/health 和 /api/control/runtime。", "已部署", "部分硬件离线时应明确降级。", "关键子系统状态可定位故障。")

    add("前端PWA", "产品", "/home 心屿产品界面", "整合情绪、专注、日记、趋势、健康和陪伴。", "启动 FastAPI；访问 /home；按导航逐页操作。", "已部署", "部分数据为浏览器本地状态。", "各页可用，实时状态能刷新。")
    add("前端PWA", "调试", "/control 调试控制台", "展示视频、FSM、trace、遥测并提供控制操作。", "访问 /control；连接 WebSocket 和控制运行时。", "已部署", "开发界面，不面向普通用户。", "状态、叠层和按钮操作一致。")
    add("前端PWA", "预览", "Page 2 独立预览", "新版视觉方案与情绪素材的静态预览。", "打开 dashboard/page2_preview/index.html；检查响应式截图。", "部分部署", "尚未替换 /home 主路由。", "桌面/移动预览无重叠和缺图。")
    add("前端PWA", "PWA", "manifest 与 Service Worker", "提供安装信息、应用壳缓存和通知点击。", "通过 HTTPS/localhost 打开；注册 SW；检查缓存和安装。", "已部署", "安全上下文要求；缓存不覆盖实时 API。", "SW 激活、静态壳可缓存、实时请求不被缓存。")
    add("前端PWA", "通知", "护眼通知", "20 分钟护眼计时到期发送 eye。", "开启 PWA 提醒和护眼计时；等待或缩短测试计时。", "已部署", "页面驱动；普通冷却 30 分钟。", "到期通知出现，点击打开健康页。")
    add("前端PWA", "通知", "久坐通知", "45 分钟久坐计时到期发送 sit。", "开启提醒和久坐计时；保持页面运行。", "已部署", "不判断真实站立活动。", "到期通知与拉伸提示同时出现。")
    add("前端PWA", "通知", "喝水通知", "09:00-22:00 内 90 分钟未记录且未达目标时发送。", "设置杯数低于目标；调整 last_at 用于测试；刷新状态。", "已部署", "依赖浏览器时间和 localStorage。", "达标/夜间不发，满足规则时只发一次。")
    add("前端PWA", "通知", "疲劳通知", "眼部/向下 gaze 异常且专注低于 60 持续 5 分钟。", "提供受控状态序列；保持 5 分钟；观察 fatigue。", "已部署", "不是医疗告警；受感知误差影响。", "不足 5 分钟不发，持续满足后发送。")
    add("前端PWA", "通知", "低专注通知", "专注记录开启且 score<50 持续 10 分钟。", "开启专注记录；注入低分状态；保持 10 分钟。", "已部署", "只在用户开启专注记录后生效。", "未开启/不足 10 分钟不发。")
    add("前端PWA", "通知", "情绪关心通知", "消费 proactive_intervention，使用 60 分钟冷却。", "触发后端策略；开启通知；观察 emotion_care。", "已部署", "不包含敏感正文；安静时段不发。", "只发送温和文本，冷却内去重。")

    add("LLM", "对话", "DeepSeek 陪伴对话", "结合当前状态、日记和用户问题生成回复。", "设置 DEEPSEEK_API_KEY；启动；POST /api/chat。", "已部署（资源待补）", "外部网络、费用和隐私需管理；有本地 fallback。", "有 key 时返回模型回复，无 key 时明确降级。")
    add("LLM", "反思", "日记反思 /api/reflect", "生成日记条目和陪伴式回应。", "准备日记/情绪上下文；POST /api/reflect；写回前端。", "已部署（资源待补）", "LLM 内容需避免诊断和强建议。", "接口成功或降级都不丢失用户日记。")
    add("LLM", "会议", "会议摘要 LLM", "根据转写文本生成摘要。", "完成 ASR；配置可选 key；POST /api/meeting/summarize。", "已部署（资源待补）", "转写误差会传递到摘要。", "摘要与原始转写大意一致。")

    add("扩展规划", "语音", "TTS 语音输出", "将陪伴文本转换为语音。", "选 Piper/MeloTTS；增加播放器和打断状态；处理回声/音量。", "未部署", "当前明确不做 TTS；会与录音争用。", "未来验收需覆盖打断、静音和回声。")
    add("扩展规划", "会议", "说话人分离", "输出 Speaker A/B/C 级发言区分。", "接 pyannote 或 embedding；与 VAD 片段对齐；再接摘要。", "未部署", "模型重、中文环境和多人重叠语音复杂。", "未来需用标注会议集评估 DER。")
    add("扩展规划", "隐私", "声纹-人脸绑定", "将说话人身份与人脸/姓名关联。", "设计同意与注册；存储特征；融合 DOA/轨迹；提供撤销。", "未部署", "生物特征隐私和错绑风险高。", "未来需人工确认和可撤销。")
    add("扩展规划", "安全", "跌倒/姿态异常", "基于姿态关键点和持续帧检测异常。", "新增 fall_detector；规则校准；二次确认；再考虑通知。", "未部署（暂不计划）", "当前不做安防；单摄像头误报风险。", "暂不验收。")
    add("扩展规划", "视觉", "VLM 场景理解", "对手动快照做自然语言场景问答。", "选本地/云 VLM；新增 scene/ask；只允许手动调用。", "未部署", "高延迟、算力、幻觉和隐私。", "未来不得驱动实时控制或安全判断。")
    add("扩展规划", "语音", "自定义唤醒词与常驻助手", "持续监听、唤醒、录音、STT、LLM、TTS 全链路。", "先按钮说话；再固定唤醒词；最后训练自定义词和打断状态机。", "部分部署", "当前只有 wake_word 基础模块。", "未来需噪声/误唤醒/设备争用长测。")
    add("扩展规划", "音频", "声音事件检测", "识别异响、玻璃破碎等声音事件。", "引入 YAMNet；治理类别；联动 DOA；加入去重确认。", "未部署（暂不计划）", "当前不做异响/安防提醒。", "暂不验收。")
    add("扩展规划", "集成", "MQTT / Home Assistant", "将状态发布为智能家居实体。", "部署 broker；设计 topic/鉴权/限流；实现 discovery。", "未部署（暂不计划）", "当前计划明确排除 MQTT。", "暂不验收。")
    add("扩展规划", "专注", "停留时长专注指标", "将 face on-screen dwell 作为辅助指标。", "增加滑动窗口；处理离开/遮挡；融合到 attention。", "未部署", "需防止把在场等同于专注。", "未来需与人工标注对比。")
    add("扩展规划", "隐私", "事件自动截图", "情绪突变等事件自动保存 JPEG。", "复用 /api/snapshot；加用户开关、冷却、目录和清理策略。", "部分部署", "现有只有手动快照；需隐私与生命周期。", "未来需验证不重复、不泄露、可删除。")
    add("扩展规划", "音频", "噪声抑制", "在 VAD/ASR 前降低环境噪声。", "选择 RNNoise/WebRTC NS；接入音频链；比较 VAD/ASR。", "未部署", "会增加 CPU 和音频延迟。", "未来用同一噪声集比较准确率。")
    add("扩展规划", "统计", "跨帧人数统计", "用滑动窗口输出稳定人数。", "复用轨迹；增加窗口/去抖；输出统计状态。", "部分部署", "当前只有帧级人数与轨迹。", "未来短遮挡不应造成数量跳变。")
    add("扩展规划", "视频", "事件视频录像", "围绕事件保存前后视频。", "增加环形缓冲、编码、配额、清理和隐私开关。", "未部署（暂不计划）", "存储和隐私成本高，偏监控。", "暂不验收。")
    add("扩展规划", "视频", "RTSP 输出", "对外提供标准实时视频流。", "部署 RTSP server；转封装/转码；鉴权和网络配置。", "未部署（暂不计划）", "偏监控基础设施。", "暂不验收。")
    add("扩展规划", "分析", "热力图", "统计画面区域停留与活动分布。", "设计匿名轨迹聚合；保存统计；增加可视化。", "未部署（暂不计划）", "与健康陪伴叙事契合度低。", "暂不验收。")
    add("扩展规划", "安防", "区域入侵检测", "在设定区域内检测进入事件。", "增加 ROI、持续帧规则、确认和通知。", "未部署（暂不计划）", "属于安防能力，当前排除。", "暂不验收。")
    add("扩展规划", "控制", "手势控制云台", "用手势驱动云台动作。", "必须重新设计权限、确认、SafetyLayer 和误触发测试。", "未部署（暂不计划）", "A 版明确禁止控制云台。", "暂不验收。")
    add("扩展规划", "通知", "ntfy / Telegram 外部推送", "服务器向外部消息服务发送提醒。", "设计账号/密钥、脱敏、重试、限流和退订。", "未部署（暂不计划）", "第一版仅 PWA 本地通知。", "暂不验收。")
    return rows


def build_catalog():
    doc = Document()
    configure_doc(doc, "compact")
    set_running_furniture(doc.sections[0], "心屿功能部署清单", "代码基准 e607758")

    for _ in range(5):
        doc.add_paragraph()
    add_title(doc, "心屿全功能部署清单", "已部署、资源待补、部分部署与规划功能统一台账", "XINYU SYSTEM REFERENCE", True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基准日期：2026-06-30  |  代码基准：origin/main / e607758")
    set_run_font(r, size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(76)
    r = p.add_run("验证结果列由项目负责人在真实设备、浏览器和模型环境中填写")
    set_run_font(r, size=10, italic=True, color=DARK_BLUE)

    doc.add_page_break()
    doc.add_heading("使用说明", level=1)
    add_callout(doc, "状态口径", "“已部署”仅表示代码已接入主链路，不代表真实硬件、浏览器或模型效果已经通过验收。")
    for text in [
        "功能范围取当前代码主链路、EXISTING_FEATURES.md、FEATURE_ROADMAP.md 与 MISSING_FEATURES_DEPLOYMENT_ASSESSMENT.md 的并集并去重。",
        "状态只使用：已部署、已部署（资源待补）、部分部署、未部署、未部署（暂不计划）。",
        "“验证标准”给出建议验收条件；“验证结果”整列留空，由项目负责人填写通过、未通过、日期和备注。",
        "手势识别框架已接入，但 models/gesture_recognizer.task 当前缺失；主动情绪干预、注视估计和 PWA 本地通知已接入。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    doc.add_heading("当前关键边界", level=2)
    p = doc.add_paragraph()
    p.add_run("控制：").bold = True
    p.add_run("main_phase3.py 是唯一真实硬件控制进程；手势、gaze、主动关心和通知不进入云台控制。")
    p = doc.add_paragraph()
    p.add_run("通知：").bold = True
    p.add_run("第一版为浏览器本地通知，不包含 ntfy、Telegram、MQTT，也不发送敏感正文。")
    p = doc.add_paragraph()
    p.add_run("产品：").bold = True
    p.add_run("跌倒、异响、安防、云台手势控制、TTS 和 MQTT 均不在当前阶段计划。")

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.45)
    sec.right_margin = Inches(0.45)
    sec.header_distance = Inches(0.25)
    sec.footer_distance = Inches(0.25)
    set_running_furniture(sec, "心屿功能部署清单", "功能表 | 验证结果待填写")

    headers = ["编号", "模块", "功能", "功能说明", "具体部署流程", "当前部署状态", "依赖与限制", "验证标准", "验证结果"]
    widths = [480, 850, 1250, 1750, 2300, 1150, 1900, 1800, 920]
    grouped = {}
    for row in catalog_rows():
        grouped.setdefault(row[0], []).append(row[1:])
    number = 1
    for group_idx, (group, records) in enumerate(grouped.items()):
        if group_idx:
            doc.add_page_break()
        h = doc.add_heading(group, level=1)
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(6)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        set_table_geometry(table, widths, 120)
        set_repeat_table_header(table.rows[0])
        for i, value in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_shading(cell, LIGHT_BLUE)
            set_cell_margins(cell, 90, 80, 90, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=7.2, bold=True, color=DARK_BLUE)
        for module, name, desc, deploy, status, limits, verify in records:
            row = table.add_row()
            prevent_row_split(row)
            values = [str(number), module, name, desc, deploy, status, limits, verify, ""]
            for i, value in enumerate(values):
                cell = row.cells[i]
                set_cell_margins(cell, 80, 80, 80, 80)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if status == "已部署":
                    status_fill = PALE_GREEN
                elif "资源待补" in status or status == "部分部署":
                    status_fill = PALE_GOLD
                elif "暂不计划" in status:
                    status_fill = PALE_RED
                else:
                    status_fill = LIGHT_GRAY
                if i == 5:
                    set_cell_shading(cell, status_fill)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 5, 8) else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                r = p.add_run(value)
                set_run_font(r, size=6.7, bold=(i in (2, 5)))
            number += 1
    doc.core_properties.title = "心屿全功能部署清单"
    doc.core_properties.subject = "全功能、部署流程、状态与验证标准"
    doc.core_properties.author = "心屿项目组"
    doc.save(CATALOG)


def add_memo_metadata(doc, rows):
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + "：")
        set_run_font(r, size=10.5, bold=True, color=INK)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)


def add_fact_block(doc, label, fact, result):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(fact)
    set_run_font(r, size=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("工作结果（归纳）：")
    set_run_font(r, size=10.5, bold=True, color=MUTED)
    r = p.add_run(result)
    set_run_font(r, size=10.5)


def _feature_table(doc, headers, widths, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths, 120)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell, 90, 80, 90, 80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color=DARK_BLUE)
    for data_row in rows:
        row = table.add_row()
        prevent_row_split(row)
        status = data_row[-1]
        if status == "已部署":
            status_fill = PALE_GREEN
        elif "资源待补" in status or status == "部分部署":
            status_fill = PALE_GOLD
        else:
            status_fill = LIGHT_GRAY
        n = len(data_row)
        for i, value in enumerate(data_row):
            cell = row.cells[i]
            set_cell_margins(cell, 80, 80, 80, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == n - 1:
                set_cell_shading(cell, status_fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, n - 1) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, size=8.7, bold=(i in (0, n - 1)))


def build_report():
    doc = Document()
    configure_doc(doc, "business")
    set_running_furniture(doc.sections[0], "心屿项目周报", "2026-06-25 至 2026-06-30")

    add_title(doc, "心屿项目工作周报", "多模态感知、控制闭环、产品界面与健康陪伴能力建设", "WEEKLY STATUS REPORT", False)
    add_memo_metadata(doc, [
        ("报告周期", "2026 年 6 月 25 日至 6 月 30 日"),
        ("代码范围", "origin/main，截止 e607758"),
        ("提交概况", "16 次提交；提交标题多为 init，本报告以实际 diff 为证据"),
        ("状态", "阶段性工作总结 / 待真实设备与浏览器补充验收"),
    ])
    doc.add_paragraph()
    add_callout(doc, "本周结论", "项目从初始多模态原型推进到具备单控制平面、硬件安全闭环、产品化首页、会议链路和健康陪伴感知的可运行系统。最新一轮已接入主动情绪干预、注视趋势和 PWA 本地提醒；手势框架已接入但模型资源待补。", LIGHT_BLUE)

    doc.add_heading("1. 本周概述", level=1)
    for item in [
        "完成 /home 产品界面情绪、专注、日记、趋势、健康陪伴全流程，以及 /control 调试控制台和视频跟踪叠层。",
        "打通多人声源 yaw 跟随、会议录音→转写→摘要→日记写入完整链路。",
        "完成 #4 主动情绪干预、#9 注视趋势、#10 PWA 本地提醒；#8 手势 A 版框架已接入，模型资源待补。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, size=10.8)

    doc.add_page_break()
    doc.add_heading("2. 提交与变更概况", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1400, 1100, 1800, 5060], 120)
    headers = ["日期", "提交数", "代表提交", "主要变化"]
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=9, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    data = [
        ("06-25", "3", "fea7d24, b7c9c63, 090afe4", "初始仓库、多模态模块、模型和页面；修复人脸/FastAPI；清理重复 dashboard。"),
        ("06-26", "6", "87e2710 ... 4954840", "架构/SOP、产品首页、控制台、转写/会议/LLM、设备与控制路径调整。"),
        ("06-28", "3", "dda1d04, eff382f, 8758a96", "系统总览、设备配置、EventBus、控制租约、SafetyLayer、Node-RED 和测试。"),
        ("06-30", "4", "67c7012 ... e607758", "控制闭环完善、Orchestrator v2、Page 2、情绪素材、叠层、#4/#8/#9/#10。"),
    ]
    for values in data:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=8.7, bold=(i == 1))

    doc.add_heading("3. 本周功能清单", level=1)

    doc.add_heading("3.1  /control 调试控制台功能", level=2)
    _feature_table(doc, ["功能", "用户操作", "系统响应 / 闭环结果", "状态"], [1700, 1900, 4000, 1560], [
        ("单人人脸跟踪 + 情绪/专注分析", "点击“开始”，保持人脸入画", "云台跟随人脸 yaw/pitch；页面实时显示情绪、专注和眼部指标", "已部署"),
        ("多人声源 yaw 跟随", "点击“开始”，多人正常说话", "DOA 驱动 yaw 转向说话人；ReSpeaker LED 同步显示方向", "已部署"),
        ("会议录音 → 转写 → 摘要 → 写入日记", "点击录音开始/停止，点击生成摘要", "生成会议摘要文本，可一键写入日记", "已部署（资源待补）"),
        ("手动 D-Pad 云台调试", "进入手动会话，按方向键或回中", "云台按指令移动，实时 readback 显示位置", "已部署"),
        ("视频跟踪叠层", "正常使用 /control 页面", "叠层显示检测框、FSM 状态、决策 trace 和遥测", "已部署"),
    ])

    doc.add_heading("3.2  本周新增健康陪伴功能（#4 / #8 / #9 / #10）", level=2)
    _feature_table(doc, ["编号", "功能", "用户操作", "系统响应 / 闭环结果", "状态"], [540, 1380, 1620, 3720, 1600], [
        ("#4", "主动情绪干预", "正常使用，保持人脸入画", "持续负面情绪/疲劳约 3 分钟后，/home 出现陪伴提示；已开启通知时发送情绪关心", "已部署"),
        ("#8", "手势识别陪伴 A 版", "对着摄像头做 5 种手势", "张手→“我在听”；握拳→收起提醒；点赞/点踩→本地反馈；剪刀手→生成积极瞬间草稿", "已部署（资源待补）"),
        ("#9", "注视方向估计", "正常使用，/home 查看“视线趋势”", "显示 center/left/right/down/away；以 15% 权重辅助专注评分", "已部署"),
        ("#10", "PWA 本地提醒（6 类）", "开启提醒权限并授权", "护眼 / 久坐 / 喝水 / 疲劳 / 低专注 / 情绪关心；支持安静时段和冷却去重", "已部署"),
    ])

    doc.add_heading("3.3  /home 其他已闭环功能", level=2)
    _feature_table(doc, ["功能", "用户操作", "系统响应 / 闭环结果", "状态"], [1700, 1900, 4000, 1560], [
        ("情绪日记 + LLM 反思", "写入当日心情/文字，点击保存", "本地保留；可触发 LLM 生成温和回应（无 key 时本地 fallback）", "已部署（资源待补）"),
        ("专注记录 + 周趋势", "开启专注记录，查看趋势页", "累计本次时长；趋势页汇总 7 天情绪和专注数据", "已部署"),
        ("护眼 / 久坐计时", "在健康页手动开启", "20 / 45 分钟到期显示站内提示，并尝试发送 PWA 通知", "已部署"),
        ("饮水记录 & 呼吸/拉伸引导", "点击“喝了一杯水”，开始呼吸或拉伸", "杯数更新、进度条变化；4-7-8 呼吸动画和分步拉伸提示完整展示", "已部署"),
        ("DeepSeek 陪伴对话 / 健康建议", "在 /home 发起对话或点击生成建议", "返回结合当前状态的建议或对话；无 key 时使用本地模板", "已部署（资源待补）"),
    ])

    doc.add_heading("4. 关键技术成果", level=1)
    achievements = [
        ("产品闭环", "/home 把实时感知状态、日记、趋势、建议和本地提醒连接为完整的健康陪伴体验。"),
        ("多模态感知", "视频、人脸、情绪、眼部指标、注视方向和 DOA 声源方向进入统一状态，驱动跟踪和陪伴策略。"),
        ("降级设计", "手势模型、LLM、通知权限或硬件缺失时提供 reason/fallback，不阻断主服务。"),
    ]
    for title, text in achievements:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title + "：")
        set_run_font(r, size=10.8, bold=True, color=DARK_BLUE)
        r = p.add_run(text)
        set_run_font(r, size=10.8)

    doc.add_heading("5. 测试与验证", level=1)
    add_callout(doc, "证据口径", "Git 可确认 6 月 28 日新增 tests/test_control_closure.py 与 tests/test_hardware_adapters.py，6 月 30 日继续更新。本文不把“存在测试代码”写成“所有真实硬件验收已通过”。", LIGHT_GRAY)
    for item in [
        "控制闭环测试覆盖会话、租约、控制事件、SafetyLayer 和 FastAPI 不直接控制硬件等关键边界。",
        "硬件适配测试覆盖 RecameraClient、DOA/网络输入和 bridge 相关行为，并在 6 月 30 日继续调整。",
        "#4/#8/#9/#10 已完成代码接入与降级路径；手势真实识别因模型文件缺失尚不能验收。",
        "PWA 通知仍需在目标浏览器、HTTPS/localhost、系统通知权限和移动后台策略下逐项验证。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_heading("6. 当前限制与遗留问题", level=1)
    limits = [
        "models/gesture_recognizer.task 不在仓库，手势框架状态为已部署（资源待补）。",
        "注视趋势受眼镜、遮挡、光照和分辨率影响，不能宣称精确眼动追踪。",
        "主动情绪干预需要真实用户与录制状态序列校准阈值、频率和接受度。",
        "PWA 第一版由 /home 页面调度，不是服务器离线 Web Push；移动系统可能冻结后台页面。",
        "DeepSeek、faster-whisper、reCamera、ReSpeaker 和 Node-RED 均需要外部资源或现场环境。",
        "Page 2 仍是独立预览，尚未替换 /home 主产品页面。",
        "跌倒、声音事件、MQTT、ntfy/Telegram、TTS 和手势控制云台明确不在当前阶段。",
    ]
    for item in limits:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_heading("7. 下周建议", level=1)
    suggestions = [
        "补齐 Gesture Recognizer 模型，在真实镜头下对五个手势逐项记录准确率、稳定帧和误触发。",
        "建立 #4/#9 的录制回放测试集，覆盖多人、眼镜、低光、遮挡和不同距离，并固化阈值基线。",
        "在 Chrome/Edge 与目标移动端验证六类通知、安静时段、冷却、权限拒绝和通知点击跳转。",
        "完成 reCamera + ReSpeaker + Node-RED 的端到端验收表，区分代码通过、模拟通过和真机通过。",
        "决定 Page 2 与 /home 的合并路径，避免长期维护两套产品交互。",
        "继续保持范围克制：先稳定健康陪伴闭环，再评估 TTS、VLM 或说话人分离。",
    ]
    for idx, item in enumerate(suggestions, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    doc.add_page_break()
    doc.add_heading("附录：16 次提交索引", level=1)
    commits = [
        ("06-25", "fea7d24", "仓库初始多模态基线"), ("06-25", "b7c9c63", "FastAPI / FaceTracker 修复"),
        ("06-25", "090afe4", "Dashboard 页面收敛"), ("06-26", "87e2710", "架构、SOP、设备与控制调整"),
        ("06-26", "c41cd6f", "转写、首页、路线图与会议接口"), ("06-26", "f65d087", "首页/控制台/架构继续集成"),
        ("06-26", "2bb9908", "系统总览文档与控制台"), ("06-26", "2cd8983", "设备、主进程和 FastAPI 修订"),
        ("06-26", "4954840", "硬件客户端与控制展示调整"), ("06-28", "dda1d04", "SOP、系统总览 Word 与服务修改"),
        ("06-28", "eff382f", "设备配置、EventBus、FSM/安全重构"), ("06-28", "8758a96", "会话租约、Node-RED、DOA、测试"),
        ("06-30", "67c7012", "Orchestrator v2、Page 2、叠层、控制闭环"), ("06-30", "6aa486d", "架构、FastAPI、硬件测试修复"),
        ("06-30", "08b44db", "情绪素材、Page 2、硬件与主进程"), ("06-30", "e607758", "#4/#8/#9/#10 健康陪伴能力"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1400, 1600, 6360], 120)
    for i, value in enumerate(("日期", "提交", "按实际 diff 归纳的主题")):
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, size=9, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for values in commits:
        row = table.add_row()
        prevent_row_split(row)
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=8.8, bold=(i == 1))

    doc.core_properties.title = "心屿项目工作周报 2026-06-25 至 2026-06-30"
    doc.core_properties.subject = "Git 历史与代码变化工作总结"
    doc.core_properties.author = "心屿项目组"
    doc.save(REPORT)


if __name__ == "__main__":
    DOCS.mkdir(parents=True, exist_ok=True)
    build_catalog()
    build_report()
    print(CATALOG)
    print(REPORT)

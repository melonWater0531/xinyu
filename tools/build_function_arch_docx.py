from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "功能架构模块说明.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("reCamera Multimodal 功能架构模块说明")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("从 SOP 中单独整理的功能与架构模块说明，不包含网页打开方式、API 调用、部署启动和 Debug 命令。")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("555555")


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 系统目标与当前能力", level=1)
    doc.add_paragraph(
        "本项目将 reCamera 云台摄像头、视觉模型、ReSpeaker DOA 和心屿 Web 产品界面整合为一个多模态交互系统。"
        "当前核心能力分为单人视觉陪伴模式和多人声源跟随模式，两者互斥运行，并由统一状态快照向前端输出。"
    )
    add_table(
        doc,
        ["模式", "输入", "核心处理", "云台行为", "用户侧输出"],
        [
            ["单人模式", "reCamera 视频", "人体/人脸检测、目标锁定、情绪与专注分析", "yaw + pitch 视觉追踪", "视频状态、追踪状态、情绪、专注度、日记上下文"],
            ["多人模式", "TCP DOA", "声源角度、speech 状态、新鲜度判断", "仅 yaw 声源跟随", "DOA、声源方向、目标 yaw、多人场景状态"],
        ],
        widths=[1.0, 1.15, 1.95, 1.35, 2.05],
    )

    doc.add_heading("2. 总体功能架构", level=1)
    doc.add_paragraph(
        "系统以 FastAPI 主进程为中心，把设备视频、视觉模型、DOA 输入、云台控制、状态同步和前端产品界面统一管理。"
        "reCamera 提供视频和云台能力，ReSpeaker 或其他主机提供声源方向，前端根据统一状态展示用户侧功能和开发调试信息。"
    )
    add_numbered(
        doc,
        [
            "reCamera 通过 SSCMA WebSocket 输出 JPEG 与设备检测框。",
            "FastAPI 接收视频帧后执行视觉检测、目标管理、情绪与专注分析。",
            "多人模式下，远端主机读取 ReSpeaker DOA，并通过 TCP 文本送入 NetworkDOA。",
            "FastAPI 根据当前模式选择单人视觉追踪或多人 DOA yaw-only 跟随。",
            "云台控制通过 Node-RED Socket.IO 或 dry-run 状态输出执行。",
            "前端通过统一状态快照获得设备、模式、视觉、DOA、情绪和专注等信息。",
        ],
    )

    doc.add_heading("3. 功能模块职责", level=1)
    add_table(
        doc,
        ["模块", "主要文件/组件", "职责说明"],
        [
            ["主入口与状态编排", "recamera_fastapi.py", "统一管理视频、模式、云台、DOA、模型、前端页面和状态快照。"],
            ["视频输入", "vision/video_stream.py", "连接 reCamera SSCMA WebSocket，持续维护最新 JPEG、检测框、分辨率、FPS 与连接状态。"],
            ["DOA 输入", "audio/network_doa.py, audio/doa.py", "监听 TCP 文本，解析声源角度、speech 状态、数据新鲜度和发送端状态。"],
            ["视觉目标检测", "vision/face_tracker_v2.py, vision/pose_estimator.py, vision/mediapipe_face.py", "提供人脸、人形、关键点和面部点位，用于单人搜索、对准和追踪。"],
            ["情绪与专注", "vision/emotieff_adapter.py, vision/attention_engine.py, vision/eye_metrics.py", "基于 face crop、landmarks 和眼部指标输出情绪、置信度、专注分与专注状态。"],
            ["云台控制与安全", "core/gimbal_mode_state.py, core/control_filter.py, core/safety_layer.py", "处理模式优先级、控制平滑、死区、步长限制和真实控制安全门。"],
            ["用户侧产品界面", "dashboard/home.html", "承载情绪监测、专注度监测、多人场景跟踪、情绪日记、LLM 对话和健康建议。"],
            ["开发调试界面", "dashboard/recamera_v2_live.html", "展示实时视频、检测叠加、模式状态、DOA、云台状态和调试日志。"],
        ],
        widths=[1.35, 2.55, 3.1],
    )

    doc.add_heading("4. 单人模式功能流程", level=1)
    doc.add_paragraph(
        "单人模式以视觉输入为主，目标是让云台找到并稳定跟随当前用户，同时把情绪和专注分析沉淀为产品侧状态。"
    )
    add_table(
        doc,
        ["阶段", "含义", "核心行为"],
        [
            ["初始检查/归中", "确认是否已有可用目标", "已有脸则进入追踪；有人体则进入对准；否则回到安全中心位。"],
            ["扫描找人", "没有稳定目标时扩大搜索", "yaw 围绕中心范围扫描，pitch 保持。"],
            ["人体已发现", "人脸暂不可用但有人体框", "根据人体框预测脸部区域并调整 yaw/pitch。"],
            ["完整人脸追踪", "稳定人脸目标可用", "持续使用人脸中心控制云台，并输出情绪与专注状态。"],
        ],
        widths=[1.5, 2.0, 3.5],
    )
    doc.add_heading("单人模式检测优先级", level=2)
    add_numbered(
        doc,
        [
            "FaceTrackerV2 完整人脸目标。",
            "YOLO pose 人体目标。",
            "YuNet 或关键点 fallback。",
            "无目标时回到扫描。"
        ],
    )
    doc.add_heading("单人模式控制原则", level=2)
    add_bullets(
        doc,
        [
            "yaw 控制水平居中，pitch 控制垂直脸部位置。",
            "EMA、步长限制和死区用于降低云台抖动。",
            "手动、睡眠、待机和急停优先于 AI 自动控制。",
            "情绪与专注不是云台控制安全前置条件，模型不可用时不阻塞视频和基础追踪。",
        ],
    )

    doc.add_heading("5. 多人 TCP DOA 模式功能流程", level=1)
    doc.add_paragraph(
        "多人模式不把 ReSpeaker 挂载到 WSL，而是由真实连接设备的主机读取 DOA，再以标准 TCP 文本送入 FastAPI。"
        "这让多人声源跟随避开 USBIP、HID 权限、PortAudio 映射和 WSL 重连问题。"
    )
    add_table(
        doc,
        ["步骤", "功能含义", "状态产物"],
        [
            ["远端读取 DOA", "Windows、reCamera 或 Linux 主机读取 ReSpeaker 声源方向", "角度、speech 或 has_speech"],
            ["NetworkDOA 接收", "FastAPI 监听 TCP 文本并解析角度", "doa_deg、age、packet_count、sender_connected"],
            ["新鲜度判断", "判断 DOA 是否过期，以及 speech hold 是否仍有效", "has_speech、age、last_line"],
            ["yaw-only 跟随", "把声源方向映射为云台 yaw 目标", "target_yaw、reason、command 状态"],
            ["前端展示", "用户侧显示多人场景，调试侧显示 DOA 细节", "DOA、声源方向、跟随状态"],
        ],
        widths=[1.25, 3.05, 2.7],
    )
    doc.add_heading("多人模式控制原则", level=2)
    add_bullets(
        doc,
        [
            "0° 表示正前方，正角度为右侧，负角度为左侧。",
            "基础映射为 target_yaw = 180 + signed_doa。",
            "只控制 yaw，不控制 pitch。",
            "DOA 过期或没有 speech 时不移动。",
            "默认是 doa_only，不保存音频；录音、ASR、说话人分离和会议摘要不属于当前标准链路。",
        ],
    )

    doc.add_heading("6. 前端产品功能模块", level=1)
    doc.add_paragraph(
        "前端分为用户侧产品主界面和开发调试控制台。用户侧功能集中在 /home，开发调试集中在 /v2。"
        "本节只描述功能结构，不描述页面打开方式。"
    )
    add_table(
        doc,
        ["页面/模块", "用户任务", "核心功能"],
        [
            ["首页总览", "快速了解今日状态", "设备状态、今日情绪、专注状态、当前场景、最近日记、继续对话。"],
            ["情绪监测", "开始或暂停情绪识别", "显示当前情绪、可信度、更新时间，并可写入日记上下文。"],
            ["专注度监测", "记录专注表现", "显示专注分、记录时长、趋势摘要和分心提示。"],
            ["多人场景跟踪", "在多人对话中跟随声源方向", "开始/停止多人跟踪，显示声源方向、DOA 新鲜度和跟随状态。"],
            ["情绪日记", "记录和回看状态", "今日记录、历史记录、生成总结、手动补充、进入对话。"],
            ["LLM 对话", "围绕状态获得回应", "基于情绪日记、今日状态、专注记录或用户手动输入进行对话。"],
            ["健康建议", "获得轻量行动建议", "根据情绪、专注和互动状态生成建议，并解释建议依据。"],
        ],
        widths=[1.5, 2.0, 3.5],
    )

    doc.add_heading("7. 状态同步与模式互斥", level=1)
    doc.add_paragraph(
        "后端约每 200ms 构建统一状态快照，并通过 WebSocket 或状态快照接口提供给前端。"
        "前端不直接拼凑多个底层状态，而是从统一状态中读取 gimbal、tracking_mode、video、pose、doa、sound_follow、conversation、face_tracking、face_lock、attention 和 emotion 等字段。"
    )
    add_table(
        doc,
        ["模式", "启用内容", "关闭内容", "前端主要展示"],
        [
            ["单人模式", "face tracking、视觉追踪、情绪分析、专注分析", "sound tracking", "视觉、情绪、专注、目标锁定和追踪状态。"],
            ["多人模式", "TCP DOA、sound tracking、yaw-only 跟随", "face tracking", "DOA、声源方向、目标 yaw 和多人场景状态。"],
        ],
        widths=[1.2, 2.45, 1.55, 2.3],
    )

    doc.add_heading("8. 当前边界与后续扩展点", level=1)
    add_bullets(
        doc,
        [
            "多人模式当前标准链路是 DOA-only，不保证音频文件、ASR 转写、说话人分离或会议摘要。",
            "LLM 对话可基于产品首页的情绪、专注和日记上下文进行回应；若后端不可用，前端应诚实展示 fallback 或待接入状态。",
            "情绪和专注模型失败不应阻塞视频、追踪和基本产品体验。",
            "真实云台控制应始终服从安全层、急停和模式优先级。",
        ],
    )

    doc.core_properties.title = "reCamera Multimodal 功能架构模块说明"
    doc.core_properties.subject = "功能架构模块独立说明"
    doc.core_properties.author = "Codex"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
